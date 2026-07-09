#!/usr/bin/env python3
"""
契約書レビュー結果をWord文書に適用するスクリプト。
python-docx でコメント挿入、lxml で Track Changes (変更履歴) を挿入する。

使い方:
    python apply_review.py <入力docx> <review_items.json> <出力docx>

JSON形式:
    {
      "review_items": [
        {
          "article": "第5条",
          "risk_level": "高",
          "category": "損害賠償",
          "original_text": "検索対象テキスト（部分一致）",
          "proposed_text": "変更後テキスト（空なら変更なし＝コメントのみ）",
          "comment": "先方へのコメント",
          "issue_type": "modification | missing_clause | comment_only"
        }
      ]
    }

issue_type:
  - modification: 既存テキストの変更（Track Changes + コメント）
  - missing_clause: 不足条項の追加（文書末尾にTrack Changesで挿入 + コメント）
  - comment_only: コメントのみ（変更なし）
"""

import json
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


AUTHOR = "mijica レビュー"
INITIALS = "MJ"


def find_paragraph_by_text(doc, search_text):
    """段落からテキストを部分一致で検索し、最初にマッチした段落を返す。"""
    for p in doc.paragraphs:
        if search_text in p.text:
            return p
    return None


def find_run_containing_text(paragraph, search_text):
    """段落内のrunからテキストを含むrunを探す。
    Word文書ではテキストが複数runに分割されることがあるため、
    段落全体のテキストで検索し、該当するrun群を返す。
    """
    full_text = paragraph.text
    start_idx = full_text.find(search_text)
    if start_idx == -1:
        return None, -1, -1
    end_idx = start_idx + len(search_text)
    return paragraph.runs, start_idx, end_idx


def add_track_change_to_paragraph(paragraph, original_text, proposed_text, revision_id):
    """段落内のテキストにTrack Changes（変更履歴）を適用する。

    original_textを含むrunを見つけ、削除マーク(w:del)と挿入マーク(w:ins)を追加する。
    Wordの内部XMLではテキストが複数runに分割されることがあるため、
    段落全体で検索し、該当部分を置換する。
    """
    p_elem = paragraph._element
    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    # 段落内の全テキストを結合して検索位置を特定
    full_text = paragraph.text
    match_start = full_text.find(original_text)
    if match_start == -1:
        return False

    match_end = match_start + len(original_text)

    # 各runの文字範囲を計算
    runs = list(paragraph.runs)
    run_ranges = []
    pos = 0
    for run in runs:
        run_text = run.text or ""
        run_ranges.append((pos, pos + len(run_text), run))
        pos += len(run_text)

    # 該当するrunを特定（部分的にかかるrunも含む）
    affected_runs = []
    for r_start, r_end, run in run_ranges:
        if r_end > match_start and r_start < match_end:
            affected_runs.append((r_start, r_end, run))

    if not affected_runs:
        return False

    # 最初の該当runの挿入位置をrun削除前に記録
    first_run_elem = affected_runs[0][2]._element
    insert_idx = list(p_elem).index(first_run_elem)

    # 元のrunの書式を保持（最初の該当runから取得）
    rpr = first_run_elem.find(qn("w:rPr"))
    rpr_copy = deepcopy(rpr) if rpr is not None else None

    # before/after テキストを計算
    first_r_start = affected_runs[0][0]
    first_run_text = affected_runs[0][2].text or ""
    overlap_start_in_first = match_start - first_r_start
    before_text = first_run_text[:overlap_start_in_first] if overlap_start_in_first > 0 else ""

    last_r_start = affected_runs[-1][0]
    last_run_text = affected_runs[-1][2].text or ""
    overlap_end_in_last = match_end - last_r_start
    after_text = last_run_text[overlap_end_in_last:] if overlap_end_in_last < len(last_run_text) else ""

    # 該当runを全て削除
    for _, _, run in affected_runs:
        p_elem.remove(run._element)

    # 挿入位置に順番に要素を配置
    current_idx = insert_idx

    # マッチ前のテキストがあれば通常runとして挿入
    if before_text:
        pre_run = etree.Element(qn("w:r"))
        if rpr_copy is not None:
            pre_run.insert(0, deepcopy(rpr_copy))
        pre_t = etree.SubElement(pre_run, qn("w:t"))
        pre_t.set(qn("xml:space"), "preserve")
        pre_t.text = before_text
        p_elem.insert(current_idx, pre_run)
        current_idx += 1

    # <w:del>（削除マーク）
    del_elem = etree.Element(qn("w:del"))
    del_elem.set(qn("w:id"), str(revision_id))
    del_elem.set(qn("w:author"), AUTHOR)
    del_elem.set(qn("w:date"), now)
    del_run = etree.SubElement(del_elem, qn("w:r"))
    if rpr_copy is not None:
        del_run.insert(0, deepcopy(rpr_copy))
    del_text = etree.SubElement(del_run, qn("w:delText"))
    del_text.set(qn("xml:space"), "preserve")
    del_text.text = original_text
    p_elem.insert(current_idx, del_elem)
    current_idx += 1

    # <w:ins>（挿入マーク）
    ins_elem = etree.Element(qn("w:ins"))
    ins_elem.set(qn("w:id"), str(revision_id + 1))
    ins_elem.set(qn("w:author"), AUTHOR)
    ins_elem.set(qn("w:date"), now)
    ins_run = etree.SubElement(ins_elem, qn("w:r"))
    if rpr_copy is not None:
        ins_run.insert(0, deepcopy(rpr_copy))
    ins_text = etree.SubElement(ins_run, qn("w:t"))
    ins_text.set(qn("xml:space"), "preserve")
    ins_text.text = proposed_text
    p_elem.insert(current_idx, ins_elem)
    current_idx += 1

    # マッチ後のテキストがあれば通常runとして挿入
    if after_text:
        post_run = etree.Element(qn("w:r"))
        if rpr_copy is not None:
            post_run.insert(0, deepcopy(rpr_copy))
        post_t = etree.SubElement(post_run, qn("w:t"))
        post_t.set(qn("xml:space"), "preserve")
        post_t.text = after_text
        p_elem.insert(current_idx, post_run)

    return True


def apply_direct_replacement(paragraph, original_text, replacement_text):
    """段落内の original_text を replacement_text に「直接」置換する。

    add_track_change_to_paragraph と同じ run 特定ロジックを使うが、
    Track Changes（w:del/w:ins）ではなく**通常の run** に差し替える（提案モードではない確定記入）。
    元 run の書式（w:rPr）は保持する。成功で True。
    """
    p_elem = paragraph._element
    full_text = paragraph.text
    match_start = full_text.find(original_text)
    if match_start == -1:
        return False
    match_end = match_start + len(original_text)

    runs = list(paragraph.runs)
    run_ranges = []
    pos = 0
    for run in runs:
        run_text = run.text or ""
        run_ranges.append((pos, pos + len(run_text), run))
        pos += len(run_text)

    affected_runs = [(s, e, r) for s, e, r in run_ranges if e > match_start and s < match_end]
    if not affected_runs:
        return False

    first_run_elem = affected_runs[0][2]._element
    insert_idx = list(p_elem).index(first_run_elem)
    rpr = first_run_elem.find(qn("w:rPr"))
    rpr_copy = deepcopy(rpr) if rpr is not None else None

    first_r_start = affected_runs[0][0]
    first_run_text = affected_runs[0][2].text or ""
    before_text = first_run_text[: match_start - first_r_start] if match_start - first_r_start > 0 else ""

    last_r_start = affected_runs[-1][0]
    last_run_text = affected_runs[-1][2].text or ""
    overlap_end_in_last = match_end - last_r_start
    after_text = last_run_text[overlap_end_in_last:] if overlap_end_in_last < len(last_run_text) else ""

    for _, _, run in affected_runs:
        p_elem.remove(run._element)

    current_idx = insert_idx
    for piece in (before_text, replacement_text, after_text):
        if not piece:
            continue
        r = etree.Element(qn("w:r"))
        if rpr_copy is not None:
            r.insert(0, deepcopy(rpr_copy))
        t = etree.SubElement(r, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = piece
        p_elem.insert(current_idx, r)
        current_idx += 1
    return True


def insert_plain_paragraphs_after(paragraph, lines):
    """指定段落の直後に、書式（pPr）を引き継いだ通常段落を順に挿入する。"""
    ppr = paragraph._element.find(qn("w:pPr"))
    prev_elem = paragraph._element
    for line in lines:
        new_p = etree.Element(qn("w:p"))
        if ppr is not None:
            new_p.append(deepcopy(ppr))
        r = etree.SubElement(new_p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = line
        prev_elem.addnext(new_p)
        prev_elem = new_p


def find_signature_block_index(doc):
    """文書末尾の署名ブロックの開始位置（段落インデックス）を返す。

    署名ブロックが見つからない場合は None を返す。
    末尾から遡って署名パターンを検出し、署名ブロックの先頭段落を特定する。
    """
    import re

    # 署名ブロックの典型パターン（条項本文と誤検出しないよう、
    # 「第○条」を含む段落はスキップする）
    signature_patterns = [
        re.compile(r"[（(](住所|社名|法人名|代表者|氏名)[)）]"),  # （住所）（社名）等
        re.compile(r"本契約.*(証として|成立).*本書"),  # 本契約成立の証として本書を...
        re.compile(r"(記名押印|署名捺印).*保管"),  # 記名押印し...保管する
    ]
    # 日付パターンは単独で署名ブロックの一部とみなす
    date_pattern = re.compile(r"^[　\s]*(●|○|\d|令和|20\d{2}).*(年).*(月).*(日)[　\s]*$")
    # 条項パターン（署名ブロックから除外）
    article_pattern = re.compile(r"第[０-９\d]+条")

    paragraphs = doc.paragraphs
    total = len(paragraphs)

    # 末尾20段落を遡って署名パターンを探す
    search_start = max(0, total - 20)
    sig_start = None

    for i in range(total - 1, search_start - 1, -1):
        text = paragraphs[i].text.strip()
        if not text:
            continue
        # 条項本文は署名ブロックではない
        if article_pattern.search(text):
            continue
        for pat in signature_patterns:
            if pat.search(text):
                sig_start = i
                break
        if sig_start is None and date_pattern.search(text):
            sig_start = i

    if sig_start is None:
        return None

    # 署名ブロックの先頭を見つけたら、さらに遡って関連行を含める
    # （日付行、空行、「本契約...証として」等）
    while sig_start > 0:
        prev_text = paragraphs[sig_start - 1].text.strip()
        if not prev_text:
            sig_start -= 1
            continue
        if article_pattern.search(prev_text):
            break
        # 署名関連パターンまたは日付パターンにマッチすれば含める
        is_sig_related = False
        for pat in signature_patterns:
            if pat.search(prev_text):
                is_sig_related = True
                break
        if not is_sig_related and date_pattern.search(prev_text):
            is_sig_related = True
        if is_sig_related:
            sig_start -= 1
        else:
            break

    print(f"  [署名検出] 段落{sig_start}以降を署名ブロックとして検出: "
          f"'{paragraphs[sig_start].text[:40]}...'")

    return sig_start


def _insert_paragraph_before(doc, ref_paragraph):
    """指定段落の直前に新しい空段落を挿入して返す。"""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p_elem = OxmlElement("w:p")
    ref_paragraph._element.addprevious(new_p_elem)
    return Paragraph(new_p_elem, ref_paragraph._element.getparent())


def _create_ins_paragraph(doc, text, revision_id, ref_paragraph=None):
    """Track Changes挿入マーク付きの段落を作成する。

    ref_paragraphが指定された場合はその直前に挿入、
    Noneの場合は文書末尾に追加する。
    """
    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    if ref_paragraph is not None:
        p = _insert_paragraph_before(doc, ref_paragraph)
    else:
        p = doc.add_paragraph()

    p_elem = p._element

    ins_elem = etree.SubElement(p_elem, qn("w:ins"))
    ins_elem.set(qn("w:id"), str(revision_id))
    ins_elem.set(qn("w:author"), AUTHOR)
    ins_elem.set(qn("w:date"), now)

    ins_run = etree.SubElement(ins_elem, qn("w:r"))
    ins_text = etree.SubElement(ins_run, qn("w:t"))
    ins_text.set(qn("xml:space"), "preserve")
    ins_text.text = text

    return p


def add_missing_clause_header(doc, revision_id, ref_paragraph=None):
    """不足条項セクションのヘッダーを追加する。"""
    return _create_ins_paragraph(
        doc, "【以下、追加条項案】", revision_id, ref_paragraph
    )


def add_missing_clause(doc, proposed_text, comment_text, revision_id,
                       ref_paragraph=None):
    """不足条項をTrack Changes付きで追加する。

    ref_paragraphが指定された場合はその直前に挿入、
    Noneの場合は文書末尾に追加する。
    proposed_textに改行が含まれる場合は段落を分割する。
    """
    lines = proposed_text.split("\n")
    first_p = None

    for line_idx, line in enumerate(lines):
        if not line.strip():
            continue

        p = _create_ins_paragraph(
            doc, line.strip(), revision_id + line_idx * 2, ref_paragraph
        )

        if first_p is None:
            first_p = p

    # 最初の段落にコメントを付与（add_commentにはrunが必要なので通常runを追加）
    if first_p is not None:
        anchor_run_elem = etree.SubElement(first_p._element, qn("w:r"))
        anchor_t = etree.SubElement(anchor_run_elem, qn("w:t"))
        anchor_t.text = ""
        if first_p.runs:
            doc.add_comment(
                runs=first_p.runs[-1:],
                text=comment_text,
                author=AUTHOR,
                initials=INITIALS,
            )

    return first_p


def apply_review(input_path, json_path, output_path):
    """レビュー結果をWord文書に適用する。"""
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    review_items = data.get("review_items", [])
    if not review_items:
        print("レビュー項目がありません。")
        return

    doc = Document(input_path)
    revision_id = 100  # 既存のrevision IDと衝突しないよう大きめの値から開始
    results = {"success": 0, "comment_only": 0, "missing": 0, "not_found": 0, "direct": 0}

    # modification / comment_only を先に処理し、missing_clause は後でまとめて末尾に追加
    missing_items = []

    for item in review_items:
        issue_type = item.get("issue_type", "modification")
        article = item.get("article", "")
        comment_text = item.get("comment", "")
        risk_level = item.get("risk_level", "")
        original_text = item.get("original_text", "")
        proposed_text = item.get("proposed_text", "")

        # コメントは先方向けの文面のみ（リスクレベル【高】【中】や条番号などの
        # 機械的ラベルは付けない＝AIレビュー感を出さない）。risk_level/article は
        # Markdownレビュー結果の整理用にのみ使う。
        full_comment = comment_text

        if issue_type == "missing_clause":
            # 後でまとめて処理
            missing_items.append((item, full_comment))
            continue

        elif issue_type == "direct_fill":
            # 直接記入（提案モードではない確定記入・コメントなし）。
            # 甲乙の当事者名（乙＝株式会社mijica）や署名欄の自社情報を埋めるのに使う。
            # proposed_text に改行が含まれる場合: 1行目を該当箇所に直接置換し、
            # 2行目以降は直後に書式を引き継いだ段落として挿入する（署名欄の社名・代表者行など）。
            p = find_paragraph_by_text(doc, original_text)
            if p:
                lines = proposed_text.split("\n")
                ok = apply_direct_replacement(p, original_text, lines[0])
                if ok and len(lines) > 1:
                    insert_plain_paragraphs_after(p, lines[1:])
                if ok:
                    results["direct"] += 1
                    print(f"  [記入] {article}: 直接記入を適用")
                else:
                    results["not_found"] += 1
                    print(f"  [失敗] {article}: 直接記入の適用に失敗")
            else:
                results["not_found"] += 1
                print(f"  [未検出] {article}: テキスト '{original_text[:30]}...' が見つかりません")

        elif issue_type == "comment_only":
            # コメントのみ: テキストを探してコメント追加
            p = find_paragraph_by_text(doc, original_text)
            if p and p.runs:
                doc.add_comment(
                    runs=p.runs,
                    text=full_comment,
                    author=AUTHOR,
                    initials=INITIALS,
                )
                results["comment_only"] += 1
                print(f"  [コメント] {article}: コメントを追加")
            else:
                results["not_found"] += 1
                print(f"  [未検出] {article}: テキスト '{original_text[:30]}...' が見つかりません")

        else:  # modification
            # テキスト変更: Track Changes + コメント
            p = find_paragraph_by_text(doc, original_text)
            if p:
                # コメントは Track Changes の「前」に付与する。
                # Track Changes 適用後はテキストが <w:ins>/<w:del> に格納され
                # python-docx の paragraph.runs から外れて検出できなくなるため、
                # 変更前の段落 run にコメントアンカーを設定する。
                if p.runs:
                    doc.add_comment(
                        runs=p.runs,
                        text=full_comment,
                        author=AUTHOR,
                        initials=INITIALS,
                    )
                if proposed_text:
                    success = add_track_change_to_paragraph(
                        p, original_text, proposed_text, revision_id
                    )
                    revision_id += 2
                    if success:
                        print(f"  [変更] {article}: Track Changes を適用")
                    else:
                        print(f"  [失敗] {article}: Track Changes の適用に失敗")
                results["success"] += 1
            else:
                results["not_found"] += 1
                print(f"  [未検出] {article}: テキスト '{original_text[:30]}...' が見つかりません")

    # missing_clause をまとめて追加（署名ブロックがあればその前に挿入）
    if missing_items:
        sig_idx = find_signature_block_index(doc)
        if sig_idx is not None:
            ref_p = doc.paragraphs[sig_idx]
            print(f"  [挿入位置] 署名ブロック（段落{sig_idx}）の前に追加条項を挿入")
        else:
            ref_p = None
            print(f"  [挿入位置] 文書末尾に追加条項を挿入")

        add_missing_clause_header(doc, revision_id, ref_p)
        revision_id += 2
        print(f"  [ヘッダー] 追加条項案セクションを挿入")

        for item, full_comment in missing_items:
            proposed_text = item.get("proposed_text", "")
            article = item.get("article", "")
            # proposed_textの行数に応じてrevision_idを確保
            line_count = len([l for l in proposed_text.split("\n") if l.strip()])
            add_missing_clause(doc, proposed_text, full_comment, revision_id, ref_p)
            revision_id += max(line_count * 2, 2)
            results["missing"] += 1
            print(f"  [追加] {article}: 不足条項を挿入")

    # 保存
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"\n完了: {output_path}")
    print(f"  変更適用: {results['success']}件")
    print(f"  直接記入: {results['direct']}件")
    print(f"  コメントのみ: {results['comment_only']}件")
    print(f"  不足条項追加: {results['missing']}件")
    print(f"  未検出: {results['not_found']}件")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print(f"Usage: {sys.argv[0]} <input.docx> <review_items.json> <output.docx>")
        sys.exit(1)

    apply_review(sys.argv[1], sys.argv[2], sys.argv[3])