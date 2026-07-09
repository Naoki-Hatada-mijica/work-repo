#!/usr/bin/env python3
"""Google Drive のファイル URL / file ID からファイルをダウンロードし、本文テキストを抽出して stdout に出す。

スキルシート（PDF / Word / Excel / Google ネイティブ）の中身を要件○×判定の根拠に取り込むためのユーティリティ。

使い方:
    python3 fetch_drive_text.py <drive_url_or_file_id>

動作:
  - URL から file ID を抽出（/file/d/<id>/ や ?id=<id> 形式、もしくは生の file ID に対応）
  - Google Sheets/Docs/Slides は xlsx/docx/pptx にエクスポートしてからダウンロード
  - 通常ファイルはそのままダウンロード
  - .pdf / .docx / .xlsx / .txt / .md からテキスト抽出して stdout に出力
  - 失敗時は stderr にエラー、exit code 非0

依存: pip3 install pdfplumber python-docx openpyxl google-api-python-client google-auth
"""
from __future__ import annotations

import io
import re
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, __import__("os").path.expanduser("~/.claude/snippets"))

from googleapiclient.http import MediaIoBaseDownload  # noqa: E402
import google_workspace  # noqa: E402

EXPORT_MAP = {
    "application/vnd.google-apps.spreadsheet": (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xlsx",
    ),
    "application/vnd.google-apps.document": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".docx",
    ),
    "application/vnd.google-apps.presentation": (
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        ".pptx",
    ),
}


def extract_file_id(raw: str) -> str:
    raw = raw.strip()
    m = re.search(r"/file/d/([A-Za-z0-9_-]+)", raw)
    if m:
        return m.group(1)
    m = re.search(r"/d/([A-Za-z0-9_-]+)", raw)
    if m:
        return m.group(1)
    m = re.search(r"[?&]id=([A-Za-z0-9_-]+)", raw)
    if m:
        return m.group(1)
    # 生の file ID とみなす
    if re.fullmatch(r"[A-Za-z0-9_-]{20,}", raw):
        return raw
    raise ValueError(f"Drive の file ID を URL から抽出できませんでした: {raw}")


def download(file_id: str) -> Path:
    svc = google_workspace.get_drive_service()
    meta = svc.files().get(
        fileId=file_id,
        fields="id, name, mimeType",
        supportsAllDrives=True,
    ).execute()
    mime = meta["mimeType"]
    name = meta.get("name", file_id)

    if mime in EXPORT_MAP:
        export_mime, ext = EXPORT_MAP[mime]
        request = svc.files().export_media(fileId=file_id, mimeType=export_mime)
        suffix = ext
    else:
        request = svc.files().get_media(fileId=file_id, supportsAllDrives=True)
        suffix = Path(name).suffix or ".bin"

    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()

    tmp = Path(tempfile.mkstemp(suffix=suffix)[1])
    tmp.write_bytes(fh.getvalue())
    print(f"[fetch_drive_text] downloaded: {name} ({mime}) -> {tmp}", file=sys.stderr)
    return tmp


def extract_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        pdfplumber = None
    if pdfplumber is not None:
        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        joined = "\n".join(parts).strip()
        if joined:
            return joined
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        raise RuntimeError("PDF 抽出には pdfplumber または pypdf が必要です") from e
    reader = PdfReader(str(path))
    return "\n".join((p.extract_text() or "") for p in reader.pages).strip()


def extract_docx(path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    lines: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(line for line in lines if line.strip())


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook  # type: ignore

    wb = load_workbook(str(path), data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"### Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cell.strip() for cell in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".txt": extract_text,
    ".md": extract_text,
}


def extract(path: Path) -> str:
    extractor = EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        raise RuntimeError(f"未対応のフォーマット: {path.suffix}")
    return extractor(path)


def main(argv: list[str]) -> int:
    if not argv:
        print(__doc__, file=sys.stderr)
        return 1
    try:
        file_id = extract_file_id(argv[0])
        path = download(file_id)
        text = extract(path)
    except Exception as e:
        print(f"!! fetch_drive_text 失敗: {e}", file=sys.stderr)
        return 1
    if not text.strip():
        print("!! 抽出結果が空です（スキャン画像のみのPDF等の可能性）", file=sys.stderr)
        return 2
    print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
