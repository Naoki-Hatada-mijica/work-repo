#!/usr/bin/env python3
"""候補者スキルシート／案件票のテキスト抽出ユーティリティ。

対応フォーマット:
- .pdf  (pdfplumber / pypdf)
- .docx (python-docx)
- .xlsx (openpyxl)
- .txt / .md (そのまま)

使い方:
    python3 extract.py <path> [<path> ...]

出力: 各ファイルごとに区切り線付きの本文を stdout に出す。
"""

from __future__ import annotations

import sys
from pathlib import Path


def extract_pdf(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        pdfplumber = None

    if pdfplumber is not None:
        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                parts.append(text)
        joined = "\n".join(parts).strip()
        if joined:
            return joined

    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "PDF 抽出には pdfplumber または pypdf が必要です: pip3 install pdfplumber"
        ) from e

    reader = PdfReader(str(path))
    return "\n".join((p.extract_text() or "") for p in reader.pages).strip()


def extract_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "Word 抽出には python-docx が必要です: pip3 install python-docx"
        ) from e

    doc = Document(str(path))
    lines: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(line for line in lines if line.strip())


def extract_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "Excel 抽出には openpyxl が必要です: pip3 install openpyxl"
        ) from e

    wb = load_workbook(str(path), data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"### Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cell.strip() for cell in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".txt": extract_text,
    ".md": extract_text,
}


def extract(path: Path) -> str:
    suffix = path.suffix.lower()
    extractor = EXTRACTORS.get(suffix)
    if extractor is None:
        raise RuntimeError(f"未対応のフォーマット: {suffix}")
    return extractor(path)


def main(argv: list[str]) -> int:
    if not argv:
        print(__doc__, file=sys.stderr)
        return 1
    for raw in argv:
        path = Path(raw)
        if not path.exists():
            print(f"!! file not found: {path}", file=sys.stderr)
            continue
        print(f"===== FILE: {path.name} =====")
        try:
            print(extract(path))
        except Exception as e:
            print(f"!! 抽出失敗: {e}", file=sys.stderr)
        print()
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
